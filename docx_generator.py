"""
版本：20260827-UX-FIXES-ROUND17
更新內容：修正文件標題（Heading 1）沒有變成標楷體的問題——實際產生一份 docx、直接
檢查 XML 確認：Word 內建的 Heading 1/2 樣式預設帶「主題字型參照」
（w:asciiTheme／w:eastAsiaTheme 等），這幾個屬性如果跟明確指定的 w:ascii/w:eastAsia
同時存在，Word 顯示時會以主題參照為準、忽略明確指定的字型。Normal 樣式（一般內文）
本來就沒有主題參照，所以之前一直是正常的，只有標題不對。_apply_font_to_style() 現在
會移除這些主題參照屬性。

---（以下為 ROUND16 紀錄）---
PDF 電子發票改成「滿版整頁」呈現——使用者反映清楚完整的 PDF 電子發票被
縮進跟手機拍照片一樣的 3.1 吋並排剪貼簿格子裡，字會小到看不清楚。新增
_add_full_page_image()，依 item.source_filename 副檔名判斷是不是 PDF：是 PDF 的
單據不進 2 欄表格，改成每張各自佔滿一整頁的寬度（維持原始比例）；同一科目底下如果
同時有照片收據跟 PDF，照片還是照舊排並排表格，PDF 接在表格後面各自分頁呈現，兩者
互不干擾。測試見 tests/test_docx_generator.py 新增的
test_pdf_receipt_uses_full_page_width_not_grid_cell／
test_pdf_and_photo_receipts_coexist_in_same_category。

---（以下為 ROUND14 紀錄）---
更新內容：剪貼簿表格每一列加上 w:cantSplit——使用者實測發現同一格的說明文字跟圖片
會被 Word 從中間切開分到不同頁（文字留在前一頁、圖片被推到下一頁，看起來像對不上），
原因是表格列預設允許「當頁放不下時從中間分頁」。強制整列當成不可分割的單位，放不下
就整列一起移到下一頁。

---（以下為 ROUND7 紀錄）---
1. 剪貼簿項目改成依「科目」分組，組與組之間強制分頁——先前是把所有收據攤平、
   兩張一列硬排，科目數量是奇數時，最後一張會跟下一個科目的第一張擠在同一列/同一頁，
   看起來像混排。
2. 機票票根的說明文字移除原始檔名（改成單純顯示「來回機票票根」），檔名對填表人沒有
   意義；同時移除文件最後條列所有機票檔名的「來回機票票根附件」區塊，圖片已經貼在
   上面的剪貼簿裡了，這個清單只是重複列檔名，沒有實際用途。

---（以下為先前版本紀錄）---
整份文件字體統一改成標楷體——中文字型要同時設定 ascii 與 eastAsia 兩個屬性
（python-docx 的 font.name 只會設定 ascii，中文實際顯示看的是 eastAsia，只設一個中文
字型不會生效），套用在 Normal/標題樣式上，讓所有段落統一繼承，不用每個 run 各別設定。

Word 明細清單從純表格改成「剪貼簿」風格——每張收據/機票票根上面一行說明文字
（日期＋內容＋幣別金額），下面貼上收據原圖，兩張並排排版節省篇幅。PDF 電子發票用 PyMuPDF
轉成圖片後嵌入；HEIC/HEIF 用 pillow-heif 轉換；圖片本身處理失敗時（檔案損毀等）不讓整份
文件產生失敗，改成只顯示說明文字＋「圖片無法嵌入」提示。
"""
import io

import pymupdf
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image
from pillow_heif import register_heif_opener

import trip_calculations
from models import ReceiptItem, TripHeader

register_heif_opener()

_IMAGE_CELL_WIDTH = Inches(3.1)
_FONT_NAME = "標楷體"


_THEME_FONT_ATTRS = ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme")


def _apply_font_to_style(doc, style_name: str, font_name: str = _FONT_NAME):
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    # Word 內建樣式（例如 Heading 1/2）預設帶「主題字型參照」（w:asciiTheme／
    # w:eastAsiaTheme 等）——這幾個屬性如果跟我們明確指定的 w:ascii/w:eastAsia/w:hAnsi
    # 同時存在，Word 實際顯示時會以主題參照為準、忽略我們明確指定的字型（實測產生的
    # docx、直接檢查 XML 確認過），這正是標題沒有變成標楷體、其他用 Normal 樣式的內文
    # 卻正常的原因——Normal 樣式本來就沒有主題參照。移除主題參照屬性，明確指定的字型
    # 才會真正生效。
    for theme_attr in _THEME_FONT_ATTRS:
        if rfonts.get(qn(theme_attr)) is not None:
            del rfonts.attrib[qn(theme_attr)]


def _set_document_font(doc, font_name: str = _FONT_NAME):
    for style_name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
        _apply_font_to_style(doc, style_name, font_name)


def _set_run_font(run, font_name: str = _FONT_NAME):
    """雙保險：除了樣式繼承，每個直接建立的 run 也明確設定字型，避免任何情況下沒繼承到。"""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    return run


def _prepare_image(raw_bytes: bytes, filename: str) -> io.BytesIO | None:
    """把任意支援格式的收據檔案轉成可以直接塞進 Word 的圖片 stream；處理失敗回傳 None。"""
    if not raw_bytes:
        return None
    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            doc = pymupdf.open(stream=raw_bytes, filetype="pdf")
            page = doc.load_page(0)
            pix = page.get_pixmap(dpi=150)
            return io.BytesIO(pix.tobytes("png"))
        # HEIC/HEIF 已透過 register_heif_opener() 讓 Pillow 能直接開啟；
        # jpg/png/webp 等其他格式 Pillow 原生支援，統一走同一條路徑重新編碼成 PNG，
        # 避免部分瀏覽器上傳的 webp/heic 檔案讓 python-docx 內部的 Pillow 呼叫失敗。
        image = Image.open(io.BytesIO(raw_bytes))
        image = image.convert("RGB")
        buf = io.BytesIO()
        image.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception:
        return None


def _prevent_row_split_across_pages(row):
    """設定該表格列的 w:cantSplit，禁止 Word 把這一列從中間切開分成兩頁——沒設這個時，
    如果某一格的圖片太高、剩餘頁面空間放不下，Word 預設允許把「說明文字留在這一頁、
    圖片被推到下一頁」，導致同一格的說明文字跟圖片各自跑到不同頁，看起來像對不上。
    強制整列當成一個不可分割的單位，放不下就整列一起移到下一頁。"""
    trPr = row._tr.get_or_add_trPr()
    cant_split = trPr.makeelement(qn("w:cantSplit"), {})
    trPr.append(cant_split)


def _add_caption_and_image(cell, caption: str, image_stream: io.BytesIO | None):
    caption_p = cell.paragraphs[0]
    run = _set_run_font(caption_p.add_run(caption))
    run.font.size = Pt(10)
    run.bold = True

    if image_stream is not None:
        image_p = cell.add_paragraph()
        image_run = image_p.add_run()
        try:
            image_run.add_picture(image_stream, width=_IMAGE_CELL_WIDTH)
        except Exception:
            note_run = _set_run_font(cell.add_paragraph().add_run("（圖片無法嵌入）"))
            note_run.font.size = Pt(9)
    else:
        note_run = _set_run_font(cell.add_paragraph().add_run("（圖片無法嵌入）"))
        note_run.font.size = Pt(9)


def _receipt_caption(item: ReceiptItem) -> str:
    date_display = f"{item.date.month}/{item.date.day}" if item.date else ""
    amount_display = f"{item.currency}:{item.amount:.0f}" if item.amount is not None else ""
    return " ".join(part for part in [date_display, item.description, amount_display] if part)


def _usable_page_width_in(doc) -> float:
    section = doc.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 914400


def _add_full_page_image(doc, caption: str, image_stream: io.BytesIO | None):
    """PDF 電子發票本身就是清楚完整的一整頁內容（不像手機拍照片還要裁切去背），縮到跟
    照片收據一樣的 3.1 吋剪貼簿格子裡，字會小到看不清楚。改成每張 PDF 單據自己獨立佔滿
    一整頁寬度，維持原始比例，不跟其他收據並排。"""
    caption_p = doc.add_paragraph()
    run = _set_run_font(caption_p.add_run(caption))
    run.font.size = Pt(11)
    run.bold = True

    if image_stream is not None:
        image_p = doc.add_paragraph()
        image_run = image_p.add_run()
        try:
            image_run.add_picture(image_stream, width=Inches(_usable_page_width_in(doc)))
            return
        except Exception:
            pass
    note_run = _set_run_font(doc.add_paragraph().add_run("（圖片無法嵌入）"))
    note_run.font.size = Pt(9)


def generate_receipt_list_docx(
    header: TripHeader,
    receipts: list[ReceiptItem],
    flight_ticket_files: list[tuple[str, bytes]] | None = None,
) -> bytes:
    doc = Document()
    _set_document_font(doc)

    title_parts = []
    if header.trip_start and header.trip_end:
        title_parts.append(
            f"{header.trip_start.year}/{header.trip_start.month}/{header.trip_start.day}"
            f"-{header.trip_end.year}/{header.trip_end.month}/{header.trip_end.day}"
        )
    destination_label = trip_calculations.destination_plain_name(header)
    if header.destination_code and header.destination_code != "Other":
        destination_label = f"{destination_label} {header.destination_code}"
    if destination_label:
        title_parts.append(f"出差{destination_label}")
    doc.add_heading(" ".join(title_parts) or "出差單據明細清單", level=1)

    info = doc.add_paragraph()
    _set_run_font(info.add_run(
        f"員工編號：{header.employee_id}　姓名：{header.employee_name}　"
        f"單位：{header.department}　職務：{header.title}\n"
        f"出差事由：{header.trip_reason}　申請日期：{header.apply_date}"
    ))

    # 剪貼簿項目：每張收據一格（說明文字＋圖片），兩格一列（table 技巧模擬並排），
    # 依科目分組、組與組之間強制分頁，同一科目的單據才不會跟其他科目混排在同一頁/同一列。
    # 純系統自動計算的項目（雜費津貼/餐費）沒有對應的原始檔案，不列進剪貼簿。標題文字
    # 只用科目/種類名稱，不顯示原始檔名——檔名對填表人沒有意義，只留圖片跟關鍵資訊即可。
    # entries 的第三個欄位標記是否為 PDF 電子發票——PDF 本身就是清楚完整的一整頁，跟手機
    # 拍照片一樣塞進 3.1 吋的剪貼簿格子會太小看不清楚，需要另外用「滿版整頁」排版（見
    # _add_full_page_image），跟一般照片收據的並排剪貼簿分開處理。
    groups: dict[str, list[tuple[str, io.BytesIO | None, bool]]] = {}
    for item in receipts:
        if not item.raw_bytes:
            continue
        is_pdf = item.source_filename.lower().endswith(".pdf")
        groups.setdefault(item.category, []).append(
            (_receipt_caption(item), _prepare_image(item.raw_bytes, item.source_filename), is_pdf)
        )
    if flight_ticket_files:
        groups["來回機票票根"] = [
            ("來回機票票根", _prepare_image(raw_bytes, filename), filename.lower().endswith(".pdf"))
            for filename, raw_bytes in flight_ticket_files
        ]

    is_first_group = True
    for category, entries in groups.items():
        if not entries:
            continue
        if not is_first_group:
            doc.add_page_break()
        is_first_group = False

        category_p = doc.add_paragraph()
        category_run = _set_run_font(category_p.add_run(category))
        category_run.bold = True
        category_run.font.size = Pt(13)

        grid_entries = [(caption, image) for caption, image, is_pdf in entries if not is_pdf]
        pdf_entries = [(caption, image) for caption, image, is_pdf in entries if is_pdf]

        if grid_entries:
            table = doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i in range(0, len(grid_entries), 2):
                row = table.add_row()
                _prevent_row_split_across_pages(row)
                row_cells = row.cells
                _add_caption_and_image(row_cells[0], *grid_entries[i])
                if i + 1 < len(grid_entries):
                    _add_caption_and_image(row_cells[1], *grid_entries[i + 1])

        for pdf_idx, (caption, image_stream) in enumerate(pdf_entries):
            # 前面有並排剪貼簿表格、或這已經不是本科目第一張 PDF，才需要先分頁——
            # 科目底下第一項內容緊接在科目標題後面就好，不必多一次分頁留白頁。
            if grid_entries or pdf_idx > 0:
                doc.add_page_break()
            _add_full_page_image(doc, caption, image_stream)

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
