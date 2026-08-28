"""
私車公用報支頁面掛載點 —— 直接沿用 car-expense-app 既有的完整邏輯（Streamlit Cloud
不支援 Git Submodule，改用複製檔案的方式併入，跟 shared-core 的 auth.py/database.py
是同一套處理原則）。這支頁面只在最上面加了「登入身分檢查」跟「SOP 下載按鈕」兩塊，
中間到最後全部是原本 car-expense-app/app.py 的內容，沒有更動任何辨識/匯出邏輯。
"""
import base64
import datetime
import io
import json
import os
import tempfile
import time
from zoneinfo import ZoneInfo
import fitz  # PyMuPDF
import google.generativeai as genai
import openpyxl

import dept_directory
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageOps
import requests
import streamlit as st
import streamlit.components.v1 as components

import database

if not st.session_state.get("logged_in"):
    st.warning("請先登入")
    st.page_link("app.py", label="回登入頁")
    st.stop()

APP_VERSION = "20260827-DEPT-STAFF-CSV-LIST"

st.set_page_config(
    page_title=f"私車公用補助單自動化工具 ({APP_VERSION})", layout="centered"
)

conn = database.get_connection()
try:
    _sop = database.get_sop_document(conn, "car_expense")
finally:
    conn.close()
if _sop:
    st.download_button("📄 查看操作SOP", data=_sop["content"], file_name=_sop["filename"])

# 注入蘋果風格 (Apple-style) 視覺樣式 CSS
# 說明：以下僅為畫面顯示樣式調整 (字型/圓角/陰影/配色/間距)，
# 不影響任何辨識、匯出等功能邏輯。
apple_style_css = """
<style>
:root {
    --apple-bg: #f4f5f7;
    --apple-card-bg: #ffffff;
    --apple-accent: #0071e3;
    --apple-accent-dark: #005bb5;
    --apple-border: #e2e2e6;
    --apple-text: #1d1d1f;
    --apple-subtext: #6e6e73;
    --apple-font: -apple-system, BlinkMacSystemFont, "SF Pro Text",
        "PingFang TC", "Microsoft JhengHei", "Helvetica Neue", Arial,
        sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--apple-font) !important;
    color: var(--apple-text);
}

/* 整體頁面背景 */
.stApp {
    background: var(--apple-bg);
}

/* 主內容區塊變成白色圓角卡片 */
[data-testid="stMainBlockContainer"],
.main .block-container {
    background: var(--apple-card-bg);
    border-radius: 20px;
    box-shadow: 0 6px 24px rgba(0, 0, 0, 0.06);
    padding: 2.2rem 2.4rem 3rem !important;
    max-width: 760px;
    margin-top: 1.5rem;
    margin-bottom: 1.5rem;
}

/* 標題與副標文字 */
h1, h2, h3 {
    color: var(--apple-text) !important;
    letter-spacing: 0.2px;
}
.stMarkdown p {
    color: var(--apple-subtext);
}

/* 下拉選單 (部門 / 姓名) */
div[data-baseweb="select"] > div {
    border-radius: 12px !important;
    border-color: var(--apple-border) !important;
    background: #f9f9fb !important;
    box-shadow: none !important;
}

/* 文字輸入框 */
.stTextInput input {
    border-radius: 10px !important;
    border-color: var(--apple-border) !important;
    background: #f9f9fb !important;
}

/* 數字輸入框 (含 +/- 按鈕的整個容器一起做圓角，避免邊角不對稱) */
[data-testid="stNumberInputContainer"] {
    border-radius: 10px !important;
    overflow: hidden;
    border: 1px solid var(--apple-border) !important;
    background: #f9f9fb !important;
}
[data-testid="stNumberInputField"] {
    border: none !important;
    background: transparent !important;
}
.stTextInput input:focus,
[data-testid="stNumberInputContainer"]:focus-within,
div[data-baseweb="select"] > div:focus-within {
    border-color: var(--apple-accent) !important;
    box-shadow: 0 0 0 3px rgba(0, 113, 227, 0.15) !important;
}

/* 檔案上傳拖拉區塊 */
[data-testid="stFileUploaderDropzone"] {
    border-radius: 14px !important;
    border: 1.5px dashed #b8d4fb !important;
    background: #f6faff !important;
}
[data-testid="stFileUploaderDropzone"] button {
    border-radius: 20px !important;
    background: var(--apple-accent) !important;
    color: #fff !important;
    border: none !important;
    box-shadow: 0 2px 8px rgba(0, 113, 227, 0.3);
}

/* 一般按鈕 / 下載按鈕 */
.stButton > button,
.stDownloadButton > button {
    border-radius: 20px !important;
    border: none !important;
    background: var(--apple-accent) !important;
    color: #fff !important;
    font-weight: 600;
    padding: 0.5rem 1.5rem !important;
    box-shadow: 0 2px 8px rgba(0, 113, 227, 0.3);
    transition: transform 0.12s ease, background 0.12s ease;
}
.stButton > button:hover,
.stDownloadButton > button:hover {
    background: var(--apple-accent-dark) !important;
    transform: translateY(-1px);
}

/* 勾選框文字 */
.stCheckbox label p {
    color: var(--apple-text);
}

/* 提示訊息 (success / info / warning / error) */
div[data-testid="stAlert"] {
    border-radius: 14px !important;
}

/* 分隔線 */
hr {
    border-color: #ececec !important;
}
</style>
"""
st.markdown(apple_style_css, unsafe_allow_html=True)

# 注入左下角程式版本標記 CSS
version_css = f"""
<style>
.custom-version-tag {{
    position: fixed;
    bottom: 16px;
    left: 20px;
    background-color: rgba(255, 255, 255, 0.9);
    padding: 4px 10px;
    border-radius: 12px;
    box-shadow: 0px 2px 6px rgba(0, 0, 0, 0.1);
    font-size: 0.8rem;
    color: #555;
    z-index: 999999;
    pointer-events: none;
    font-family: monospace, sans-serif;
}}
.custom-version-tag code {{
    color: #2e7d32;
    background-color: #f1f8e9;
    padding: 2px 5px;
    border-radius: 4px;
}}
</style>
<div class="custom-version-tag">
    \U0001F4CC 程式版本：<code>{APP_VERSION}</code>
</div>
"""
st.markdown(version_css, unsafe_allow_html=True)

# Google 表單背景紀錄設定
FORM_RESPONSE_URL = "https://docs.google.com/forms/d/e/1FAIpQLSeWI7dFxqjMeX9H0KxbSYVETuBiTOLEqZs43T06yKdbQofNAQ/formResponse"
ENTRY_NAME_ID = "entry.505350995"
ENTRY_DEPT_ID = "entry.1840094204"


def render_block_progress_html(percentage, current_count, total_count):
    """產出 20 格黑色方塊進度條，以 5% 為最小計數單位，數字顯示於右側"""
    pct_step_5 = int(round(percentage / 5.0)) * 5
    pct_step_5 = max(0, min(100, pct_step_5))

    filled_blocks = pct_step_5 // 5

    blocks_html = ""
    for i in range(20):
        if i < filled_blocks:
            blocks_html += '<div class="block filled"></div>'
        else:
            blocks_html += '<div class="block empty"></div>'

    html_code = f"""
    <style>
    .progress-wrapper {{
        width: 100%;
        max-width: 620px;
        margin: 15px 0 25px 0;
        font-family: Arial, sans-serif;
    }}
    .status-title {{
        font-size: 0.95rem;
        font-weight: bold;
        color: #333;
        margin-bottom: 8px;
    }}
    .bar-row {{
        display: flex;
        align-items: center;
        gap: 12px;
    }}
    .progress-container {{
        flex: 1;
        display: flex;
        height: 38px;
        border: 3px solid #000;
        background-color: #fff;
        padding: 2px;
        box-sizing: border-box;
        gap: 2px;
    }}
    .block {{
        flex: 1;
        height: 100%;
        box-sizing: border-box;
        transition: background-color 0.2s ease-in-out;
    }}
    .block.filled {{
        background-color: #000;
    }}
    .block.empty {{
        background-color: #fff;
        border-right: 1px solid #ddd;
    }}
    .block.empty:last-child {{
        border-right: none;
    }}
    .pct-display {{
        font-size: 1.25rem;
        font-weight: 800;
        color: #000;
        min-width: 55px;
        text-align: right;
    }}
    </style>
    
    <div class="progress-wrapper">
        <div class="status-title">
            \U0001F504 Gemini 分析單據中 ({current_count}/{total_count} 張)
        </div>
        <div class="bar-row">
            <div class="progress-container">
                {blocks_html}
            </div>
            <div class="pct-display">
                {pct_step_5}%
            </div>
        </div>
    </div>
    """
    return html_code


def log_usage_to_google_form(name_val, dept_val):
    """背景默默傳送使用者姓名與部門紀錄至 Google 表單 (完全靜音無感)"""
    try:
        form_data = {
            ENTRY_NAME_ID: name_val if name_val else "NA",
            ENTRY_DEPT_ID: dept_val if dept_val else "NA",
        }
        headers = {
            "Content-Type": "application/x-www-form-urlencoded",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        }

        session = requests.Session()
        session.post(
            FORM_RESPONSE_URL,
            data=form_data,
            headers=headers,
            timeout=5,
            allow_redirects=True,
        )
    except Exception:
        pass


# 注入 JavaScript：強效導航演算法 + 欄位 Enter 導航
def inject_enter_and_memory_js():
    js_code = """
    <script>
    function setupInteractions() {
        const doc = window.parent.document;
        
        function attachEnterListeners() {
            const allInputs = Array.from(doc.querySelectorAll('input[type="text"], input[type="number"]'));
            allInputs.forEach((input) => {
                if (!input.dataset.enterBound) {
                    input.dataset.enterBound = "true";
                    input.addEventListener('keydown', function(e) {
                        if (e.key === 'Enter') {
                            e.preventDefault();
                            const validInputs = Array.from(doc.querySelectorAll('input[type="text"]:not([disabled]), input[type="number"]:not([disabled])'));
                            const currentIndex = validInputs.indexOf(this);
                            if (currentIndex !== -1 && currentIndex < validInputs.length - 1) {
                                const nextInput = validInputs[currentIndex + 1];
                                setTimeout(() => {
                                    nextInput.focus();
                                    if (typeof nextInput.select === 'function') {
                                        nextInput.select();
                                    }
                                }, 100);
                            }
                        }
                    });
                }
            });
        }

        attachEnterListeners();
        setInterval(attachEnterListeners, 1000);
    }
    
    setTimeout(setupInteractions, 300);
    </script>
    """
    components.html(js_code, height=0, width=0)


# 注入個人專屬署名 (置於畫面底部正中間)
def inject_custom_footer():
    avatar_candidates = ["avatar.jpg", "avatar.jpeg", "avatar.png", "avatar.JPG"]
    img_base64 = ""
    mime_type = "image/png"

    for af in avatar_candidates:
        if os.path.exists(af):
            with open(af, "rb") as img_f:
                img_base64 = base64.b64encode(img_f.read()).decode("utf-8")
                if af.lower().endswith((".jpg", ".jpeg")):
                    mime_type = "image/jpeg"
                else:
                    mime_type = "image/png"
            break

    avatar_html = (
        f'<img src="data:{mime_type};base64,{img_base64}" style="width: 36px; height: 36px; border-radius: 50%; object-fit: cover; margin-right: 8px; border: 1.5px solid #ccc; background-color: #fff;">'
        if img_base64
        else ""
    )

    footer_css = f"""
    <style>
    .custom-footer-max {{
        position: fixed;
        bottom: 16px;
        left: 50%;
        transform: translateX(-50%);
        display: flex;
        align-items: center;
        background-color: rgba(255, 255, 255, 0.95);
        padding: 4px 14px;
        border-radius: 20px;
        box-shadow: 0px 2px 8px rgba(0, 0, 0, 0.15);
        z-index: 999999;
        pointer-events: none;
    }}
    .custom-footer-text {{
        font-family: 'Comic Sans MS', cursive, sans-serif;
        font-weight: bold;
        font-style: italic;
        font-size: 0.95rem;
        color: #333333;
        white-space: nowrap;
    }}
    </style>
    <div class="custom-footer-max">
        {avatar_html}
        <span class="custom-footer-text">Design by Max</span>
    </div>
    """
    st.markdown(footer_css, unsafe_allow_html=True)


# 1. 網頁標題
st.markdown(
    "<h2 style='font-size: 1.6rem; font-weight: 700; white-space: nowrap; margin-top: -10px; text-align: center;'>\U0001F697 私車公用補助單自動化填寫工具 \U0001F697</h2>",
    unsafe_allow_html=True,
)

st.markdown(
    "上傳停車/加油發票或 **PDF 檔**，由 Gemini AI 自動辨識日期與金額，輕鬆生成報銷單！"
)

# 2. API Key 設定 (從 Streamlit Secrets 讀取)
KEY_PART1 = "AQ.Ab8RN6JNdZJgY7a7BDK67Cx"
KEY_PART2 = "W44rm-vd-bHVwIkaCS84ZPG9yww"
DEFAULT_API_KEY = KEY_PART1 + KEY_PART2

api_key = st.secrets.get("GEMINI_API_KEY", DEFAULT_API_KEY)

if not api_key:
    st.error("\u26A0\uFE0F 未偵測到有效的 API Key，請確認 Streamlit Secrets 設定。")
    st.stop()

# 假身分測試進入點（見 project_spec.md 第七節）：獨立跑這支 app 開發測試時，
# 不用每次都跑完整的入口網站登入流程，用假的員工資料快速模擬「已從入口網站登入」的
# session_state。這支 app 併入 portal-app 之後，session_state 已經有 logged_in，
# 這個區塊就不會再顯示。
if not st.session_state.get("logged_in"):
    with st.sidebar.expander("🧪 開發測試：假身分登入", expanded=False):
        st.caption("僅供獨立測試使用，正式併入入口網站後這裡會自動隱藏。")
        if st.button("使用測試身分（ETW00375 溫文福）"):
            st.session_state["logged_in"] = True
            st.session_state["employee_id"] = "ETW00375"
            st.session_state["employee_name"] = "溫文福"
            st.session_state["employee_department"] = "伺服器事業部"
            st.session_state["employee_title"] = "主任工程師"
            st.rerun()

# 3. 基本資料填寫 (左側部門 -> 右側連動姓名)
# 優先使用入口網站（portal-app）登入後帶入的身分資料（session_state 的
# logged_in/employee_department/employee_name）。這支 app 併入入口網站的
# pages/ 之後，同仁不用再手動選部門/姓名，登入身分直接帶入、確認即可。
# 沒有登入資料時（獨立測試進入點，直接 streamlit run 這支 app），維持原本
# 部門/人員清單讀取（見 dept_directory.py）或手動輸入的 fallback 行為，不拿掉。
_portal_department = st.session_state.get("employee_department")
_portal_name = st.session_state.get("employee_name")

if st.session_state.get("logged_in") and _portal_department and _portal_name:
    col1, col2 = st.columns(2)
    with col1:
        st.text_input("部門", value=_portal_department, disabled=True)
    with col2:
        st.text_input("姓名", value=_portal_name, disabled=True)
    st.caption("已依入口網站登入身分自動帶入。")

    user_dept = _portal_department
    real_user_name = _portal_name
else:
    # 部門/人員清單改為讀取 GitHub 倉庫內的 department_staff.csv（見 dept_directory.py），
    # 不再寫死於程式碼中，之後只要改清單檔案就能新增/調整部門與人員，不必改程式。
    # 清單檔案讀不到（尚未建立、格式錯誤等）時，靜默改用「自己填寫」的手動輸入模式，
    # 不會顯示錯誤訊息擋住表單使用。
    _staff_directory = dept_directory.load_dept_directory()

    if _staff_directory:
        dept_options = ["NA"] + sorted(_staff_directory.keys())

        if "user_dept_selected" not in st.session_state:
            st.session_state["user_dept_selected"] = "NA"
        if "user_name_selected" not in st.session_state:
            st.session_state["user_name_selected"] = "NA"

        col1, col2 = st.columns(2)

        # 左邊：先選部門
        with col1:
            user_dept = st.selectbox("部門", dept_options, key="user_dept_selected")

        # 根據部門選擇，決定右邊姓名的動態可選項 (暱稱 -> 中文姓名對照表也一併取得)
        dept_staff = _staff_directory.get(user_dept, [])
        name_display_options = [nickname for nickname, _ in dept_staff] or ["NA"]
        name_map = dict(dept_staff)

        # 右邊：再選姓名
        with col2:
            user_name_display = st.selectbox(
                "姓名", name_display_options, key="user_name_selected"
            )

        # 自動將選取的英文暱稱轉換還原為純中文真實姓名，供後續 Excel/Word/Google 表單使用
        real_user_name = name_map.get(user_name_display, user_name_display)
    else:
        # 清單檔案不存在或讀取失敗：改為自己手動填寫部門與姓名
        col1, col2 = st.columns(2)
        with col1:
            user_dept = st.text_input("部門", key="user_dept_selected")
        with col2:
            user_name_display = st.text_input("姓名", key="user_name_selected")

        real_user_name = user_name_display

# 4. 上傳檔案
uploaded_parking_files = st.file_uploader(
    "1. 上傳停車發票/收據（照片或 PDF 檔，可多選）",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True,
)

uploaded_gas_files = st.file_uploader(
    "2. 上傳加油發票/收據（照片或 PDF 檔，可多選）",
    type=["jpg", "jpeg", "png", "pdf"],
    accept_multiple_files=True,
)


def get_taipei_now():
    """回傳台灣時區 (Asia/Taipei) 的目前時間。

    用於檔名/表單日期等「今天日期」相關欄位，避免部署主機時區
    (例如 Streamlit Cloud 預設為 UTC) 與台灣時間相差 8 小時，
    導致半夜 0~7 點產出的檔案日期誤植為前一天。
    """
    return datetime.datetime.now(ZoneInfo("Asia/Taipei"))


def format_date_to_excel(d_str):
    d_str = str(d_str).strip()
    if len(d_str) == 8 and d_str.isdigit():
        return f"{d_str[:4]}/{d_str[4:6]}/{d_str[6:]}"
    return d_str


def set_cell_value(ws, cell_ref, value):
    if isinstance(cell_ref, str):
        cell = ws[cell_ref]
    else:
        cell = ws.cell(row=cell_ref[0], column=cell_ref[1])

    if type(cell).__name__ == "MergedCell":
        for rng in ws.merged_cells.ranges:
            if cell.coordinate in rng:
                ws.cell(row=rng.min_row, column=rng.min_col).value = value
                return
    cell.value = value


def crop_and_rotate_receipt_bytes(raw_bytes, box_2d, rotate_deg):
    try:
        image = Image.open(io.BytesIO(raw_bytes))
        image = ImageOps.exif_transpose(image)

        if box_2d and isinstance(box_2d, list) and len(box_2d) == 4:
            w, h = image.size
            ymin, xmin, ymax, xmax = box_2d

            left = max(0, int(xmin * w / 1000) - 10)
            top = max(0, int(ymin * h / 1000) - 10)
            right = min(w, int(xmax * w / 1000) + 10)
            bottom = min(h, int(ymax * h / 1000) + 10)

            if right > left and bottom > top:
                image = image.crop((left, top, right, bottom))

        if rotate_deg in [90, 180, 270]:
            if rotate_deg == 90:
                image = image.transpose(Image.ROTATE_270)
            elif rotate_deg == 180:
                image = image.transpose(Image.ROTATE_180)
            elif rotate_deg == 270:
                image = image.transpose(Image.ROTATE_90)

        out_io = io.BytesIO()
        image.save(out_io, format="PNG")
        return out_io.getvalue()
    except Exception:
        return raw_bytes


def process_single_file_with_gemini(uploaded_file, receipt_type, key):
    genai.configure(api_key=key.strip())

    candidate_models = [
        "gemini-1.5-flash",
        "gemini-1.5-pro",
        "gemini-2.0-flash",
        "gemini-2.5-flash",
    ]

    try:
        available = [
            m.name.replace("models/", "")
            for m in genai.list_models()
            if "generateContent" in m.supported_generation_methods
        ]
        for a in available:
            if a not in candidate_models:
                candidate_models.insert(0, a)
    except Exception:
        pass

    prompt = """
    你是一個財務報銷助手。請讀取這份發票/收據照片或文件，並提取以下資訊：
    1. date: 發票日期，格式請統一轉換為 YYYYMMDD（例如 20260603）。若無法取得年份，預設為當前年份。
    2. amount: 金額 (數字)。
    3. box_2d: 圖片中「收據/發票紙張本體」的範圍座標 [ymin, xmin, ymax, xmax]，數值請以 0 到 1000 之間的整數表示。請貼緊收據邊緣去背。若為 PDF 則輸出 [0, 0, 1000, 1000]。
    4. rotate: 圖片中收據文字的方向。為了讓收據變成「文字由左至右、由上至下正向讀取」的直立長條狀，請判斷需要【順時針旋轉多少度】：
       - 若文字已經正面朝上：輸出 0
       - 若文字頭朝左：輸出 90
       - 若文字倒立 (頭朝下)：輸出 180
       - 若文字頭朝右：輸出 270
    
    請直接輸出純 JSON 格式，例如：
    {"date": "20260603", "amount": 150, "box_2d": [150, 250, 850, 750], "rotate": 0}
    注意：絕對不要加上 ```json 或任何 markdown 標記。
    """

    bytes_data = uploaded_file.read()
    file_ext = uploaded_file.name.split(".")[-1].lower()

    if file_ext == "pdf":
        mime_type = "application/pdf"
    elif file_ext == "png":
        mime_type = "image/png"
    else:
        mime_type = "image/jpeg"

    content_part = {"mime_type": mime_type, "data": bytes_data}

    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content([prompt, content_part])

            raw_text = response.text.strip()
            clean_text = (
                raw_text.replace("```json", "").replace("```", "").strip()
            )

            res_json = json.loads(clean_text)

            box_2d = res_json.get("box_2d", [0, 0, 1000, 1000])
            rotate_deg = res_json.get("rotate", 0)

            if file_ext in ["jpg", "jpeg", "png"]:
                processed_bytes = crop_and_rotate_receipt_bytes(
                    bytes_data, box_2d, rotate_deg
                )
            else:
                processed_bytes = bytes_data

            res_json["raw_bytes"] = processed_bytes
            res_json["file_ext"] = file_ext
            res_json["filename"] = uploaded_file.name
            res_json["receipt_type"] = receipt_type
            return res_json
        except Exception:
            continue

    return None


def set_cell_border(cell, color="D9D9D9", size=4):
    """幫 Word 表格儲存格加上一圈細邊框線 (預設淡灰色)，
    讓報支單據的收據方格看起來像整齊的表格。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), str(size))
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)
        tc_borders.append(edge_el)
    tc_pr.append(tc_borders)


def render_receipt_image_into_paragraph(paragraph, raw_bytes, file_ext, max_width_in, max_height_in):
    """把收據圖片 (照片或 PDF 第一頁) 依「等比縮放後盡量塞滿方框」的方式
    插入指定段落，回傳是否成功插入。"""
    img_stream = None
    if file_ext == "pdf":
        pdf_doc = fitz.open(stream=raw_bytes, filetype="pdf")
        if len(pdf_doc) > 0:
            page = pdf_doc[0]
            pix = page.get_pixmap(dpi=200)
            img_stream = io.BytesIO(pix.tobytes("png"))
    else:
        img_stream = io.BytesIO(raw_bytes)

    if not img_stream:
        return False

    try:
        pil_img = Image.open(img_stream)
        w, h = pil_img.size
        aspect_ratio = h / w if w > 0 else 1.0

        # 先以「高度優先」計算寬度，若寬度超出方框上限，再改以寬度為準，
        # 確保收據圖片在圓角方框內盡量放大又不會超出邊界 (contain fit)。
        target_height = max_height_in
        target_width = (
            target_height / aspect_ratio if aspect_ratio > 0 else max_width_in
        )

        if target_width > max_width_in:
            target_width = max_width_in
            target_height = target_width * aspect_ratio

        img_stream.seek(0)
        img_run = paragraph.add_run()
        img_run.add_picture(img_stream, width=Inches(target_width))
        return True
    except Exception as e:
        paragraph.add_run(f"[圖片載入失敗: {e}]")
        return False


def render_receipt_grid(doc, section_title, items, items_per_row=2):
    """把同一類 (停車費/加油費) 的收據，以多欄表格方式整齊排版輸出，
    取代「每張收據獨自佔一整頁」的舊做法，節省紙張並讓版面更整齊。"""
    if not items:
        return

    heading_p = doc.add_paragraph()
    heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_p.paragraph_format.space_before = Pt(18)
    heading_p.paragraph_format.space_after = Pt(10)
    # 標題務必跟緊接在後面的表格同頁，避免像「加油費明細」這種標題
    # 被單獨留在前一頁最下面、內容卻跑到下一頁去的情況。
    heading_p.paragraph_format.keep_with_next = True
    heading_run = heading_p.add_run(section_title)
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.name = "標楷體"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    section = doc.sections[0]
    usable_width_in = (
        section.page_width - section.left_margin - section.right_margin
    ) / 914400
    cell_padding_in = 0.15
    max_cell_width_in = max(1.0, usable_width_in / items_per_row - cell_padding_in)
    max_cell_height_in = 3.1

    row_count = -(-len(items) // items_per_row)  # 無條件進位
    table = doc.add_table(rows=row_count, cols=items_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 禁止表格「同一列」被分頁截斷，避免同一張收據的標題跟圖片被拆到不同頁
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

    for idx, item in enumerate(items):
        r, c = divmod(idx, items_per_row)
        cell = table.cell(r, c)
        set_cell_border(cell)

        title_p = cell.paragraphs[0]
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(4)
        title_p.paragraph_format.space_after = Pt(2)
        title_run = title_p.add_run(f"{item['date']} {item['title_type']}")
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.name = "標楷體"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

        img_p = cell.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(2)
        img_p.paragraph_format.space_after = Pt(6)
        render_receipt_image_into_paragraph(
            img_p,
            item["raw_bytes"],
            item["file_ext"],
            max_cell_width_in,
            max_cell_height_in,
        )

    # 項目數不是 items_per_row 的倍數時，補上空白格的邊框，維持整體表格方正
    for idx in range(len(items), row_count * items_per_row):
        r, c = divmod(idx, items_per_row)
        set_cell_border(table.cell(r, c))


has_files = (uploaded_parking_files and len(uploaded_parking_files) > 0) or (
    uploaded_gas_files and len(uploaded_gas_files) > 0
)

if has_files and real_user_name and user_dept:
    if st.button("\U0001F916 AI 辨識單據內容"):
        # 點擊 AI 辨識時發送真實全名至背景紀錄
        log_usage_to_google_form(real_user_name, user_dept)

        parsed_parking = []
        parsed_gas = []

        # 整理所有待處理檔案清單
        file_queue = []
        if uploaded_parking_files:
            for pf in uploaded_parking_files:
                file_queue.append((pf, "parking"))
        if uploaded_gas_files:
            for gf in uploaded_gas_files:
                file_queue.append((gf, "gas"))

        total_files = len(file_queue)

        # 動態進度條容器
        progress_placeholder = st.empty()

        # 初始畫面 0%
        progress_placeholder.markdown(
            render_block_progress_html(0, 0, total_files),
            unsafe_allow_html=True,
        )

        for current_idx, (f, r_type) in enumerate(file_queue):
            res = process_single_file_with_gemini(f, r_type, api_key)
            if res:
                if r_type == "parking":
                    parsed_parking.append(res)
                else:
                    parsed_gas.append(res)

            # 更新進度條趴數與 20 格方塊
            done_count = current_idx + 1
            current_pct = int((done_count / total_files) * 100)
            progress_placeholder.markdown(
                render_block_progress_html(
                    current_pct, done_count, total_files
                ),
                unsafe_allow_html=True,
            )

        # 完成後稍微停留讓使用者看見 100% 填滿狀態
        time.sleep(0.5)
        progress_placeholder.empty()

        parsed_parking.sort(key=lambda x: x["date"])
        parsed_gas.sort(key=lambda x: x["date"])

        st.session_state["parsed_parking"] = parsed_parking
        st.session_state["parsed_gas"] = parsed_gas

        tot_count = len(parsed_parking) + len(parsed_gas)
        st.success(
            f"成功辨識 {tot_count} 筆單據資料（停車: {len(parsed_parking)} 筆，加油: {len(parsed_gas)} 筆）！"
        )

# 5. 補充欄位與匯出 Excel / Word
if "parsed_parking" in st.session_state or "parsed_gas" in st.session_state:
    parking_receipts = st.session_state.get("parsed_parking", [])
    gas_receipts = st.session_state.get("parsed_gas", [])

    st.subheader("\U0001F4DD 補充填寫報銷明細 (私車公用 Excel)")

    if len(parking_receipts) > 0:
        chk_col1, chk_col2, chk_col3, chk_col4 = st.columns(4)
        same_loc = chk_col1.checkbox("所有【地點】相同")
        same_km = chk_col2.checkbox("所有【公里數】相同")
        same_toll = chk_col3.checkbox("所有【過路費】相同")
        same_reason = chk_col4.checkbox("所有【事由】相同")

        st.markdown("---")
        details = []

        first_loc, first_km, first_toll, first_reason = "", 0, 0, ""

        for idx, r in enumerate(parking_receipts):
            formatted_date = format_date_to_excel(r["date"])
            st.markdown(
                f"**停車單據 {idx+1} (日期: {formatted_date} / 金額: {r['amount']}元)**"
            )
            c1, c2, c3, c4 = st.columns(4)

            # 1. 地點欄位
            if idx == 0:
                loc_val = c1.text_input(
                    f"地點 #{idx+1}",
                    value="",
                    placeholder="例如：客戶端",
                    key=f"loc_{idx}",
                )
                first_loc = loc_val
            else:
                if same_loc:
                    loc_val = first_loc
                    c1.text_input(
                        f"地點 #{idx+1}",
                        value=first_loc,
                        disabled=True,
                        key=f"dis_loc_{idx}",
                    )
                else:
                    loc_val = c1.text_input(
                        f"地點 #{idx+1}",
                        value="",
                        placeholder="例如：客戶端",
                        key=f"loc_{idx}",
                    )

            # 2. 公里數欄位
            if idx == 0:
                km_val = int(
                    c2.number_input(
                        f"公里數 #{idx+1}", value=0, step=1, key=f"km_{idx}"
                    )
                )
                first_km = km_val
            else:
                if same_km:
                    km_val = int(first_km)
                    c2.number_input(
                        f"公里數 #{idx+1}",
                        value=int(first_km),
                        disabled=True,
                        key=f"dis_km_{idx}",
                    )
                else:
                    km_val = int(
                        c2.number_input(
                            f"公里數 #{idx+1}", value=0, step=1, key=f"km_{idx}"
                        )
                    )

            # 3. 回數票/過路費欄位
            if idx == 0:
                toll_val = int(
                    c3.number_input(
                        f"回數票 #{idx+1}", value=0, step=1, key=f"toll_{idx}"
                    )
                )
                first_toll = toll_val
            else:
                if same_toll:
                    toll_val = int(first_toll)
                    c3.number_input(
                        f"回數票 #{idx+1}",
                        value=int(first_toll),
                        disabled=True,
                        key=f"dis_toll_{idx}",
                    )
                else:
                    toll_val = int(
                        c3.number_input(
                            f"回數票 #{idx+1}", value=0, step=1, key=f"toll_{idx}"
                        )
                    )

            # 4. 事由欄位
            if idx == 0:
                reason_val = c4.text_input(
                    f"事由 #{idx+1}",
                    value="",
                    placeholder="例如：拜訪客戶",
                    key=f"reason_{idx}",
                )
                first_reason = reason_val
            else:
                if same_reason:
                    reason_val = first_reason
                    c4.text_input(
                        f"事由 #{idx+1}",
                        value=first_reason,
                        disabled=True,
                        key=f"dis_reason_{idx}",
                    )
                else:
                    reason_val = c4.text_input(
                        f"事由 #{idx+1}",
                        value="",
                        placeholder="例如：拜訪客戶",
                        key=f"reason_{idx}",
                    )

            details.append(
                {
                    "date": formatted_date,
                    "location": loc_val,
                    "km": km_val,
                    "parking": int(round(float(r["amount"]))),
                    "toll": toll_val,
                    "reason": reason_val,
                    "raw_bytes": r["raw_bytes"],
                    "file_ext": r["file_ext"],
                }
            )
    else:
        details = []
        st.info(
            "\U0001F4A1 目前未上傳停車發票，匯出 Excel 功能將暫跳過，但 Word 憑證仍可正常整合產出！"
        )

    st.markdown(" ")
    btn_col1, btn_col2 = st.columns(2)

    # 產出 Excel
    with btn_col1:
        if st.button("\U0001F680 產出 Excel 報銷檔案"):
            if len(details) == 0:
                st.warning("\u26A0\uFE0F 請先上傳並填寫至少一筆停車發票明細！")
            else:
                # 點擊產出 Excel 時發送真實全名至背景紀錄
                log_usage_to_google_form(real_user_name, user_dept)

                template_xlsx = "私車公用補助申請單.xlsx"

                if not os.path.exists(template_xlsx):
                    st.error(
                        "系統找不到範本『私車公用補助申請單.xlsx』！請確認檔名是否包含 .xlsx"
                    )
                else:
                    wb = openpyxl.load_workbook(template_xlsx)

                    ws1 = wb.worksheets[0]
                    # 填入純中文真實全名
                    set_cell_value(ws1, "B3", real_user_name)
                    set_cell_value(ws1, "E3", user_dept)

                    for i, item in enumerate(details):
                        row_num = 5 + i
                        set_cell_value(ws1, (row_num, 1), item["date"])
                        set_cell_value(ws1, (row_num, 2), item["location"])
                        set_cell_value(ws1, (row_num, 4), item["km"])
                        set_cell_value(ws1, (row_num, 5), item["parking"])
                        set_cell_value(ws1, (row_num, 6), item["toll"])
                        set_cell_value(ws1, (row_num, 8), item["reason"])

                    ws2 = wb.worksheets[1]
                    today_str = get_taipei_now().strftime("%Y年%m月%d日")

                    set_cell_value(ws2, "G5", today_str)
                    # 填入純中文真實全名
                    set_cell_value(ws2, "C7", f"專案編號：{real_user_name}")

                    first_date = details[0]["date"]
                    last_date = details[-1]["date"]
                    set_cell_value(ws2, "A9", f"{first_date}~{last_date}交通費用")

                    tot_km = sum(item["km"] for item in details)
                    tot_parking = sum(item["parking"] for item in details)
                    tot_toll = sum(item["toll"] for item in details)
                    grand_total = (tot_km * 6) + tot_parking + tot_toll

                    set_cell_value(ws2, "G9", grand_total)
                    set_cell_value(ws2, "G17", grand_total)

                    output_date = get_taipei_now().strftime("%Y%m%d")
                    # 檔名使用純中文真實全名
                    out_filename = (
                        f"私車公用補助申請單-{real_user_name}-{output_date}.xlsx"
                    )

                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".xlsx"
                    ) as tmp:
                        wb.save(tmp.name)
                        tmp_path = tmp.name

                    st.info(
                        "\U0001F4A1 **提醒：** 檔案下載後，請記得在 **「私車公用補助單」** 與 **「支出憑單」** 頁面上方，加上 **公司抬頭** 的字樣喔！"
                    )

                    with open(tmp_path, "rb") as file:
                        st.download_button(
                            label="\U0001F4E5 下載報銷單 (Excel)",
                            data=file,
                            file_name=out_filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )

    # 產出 Word 報支單據憑證檔
    with btn_col2:
        if st.button("\U0001F4C4 產出 Word 報支單據檔"):
            # 點擊產出 Word 時發送真實全名至背景紀錄
            log_usage_to_google_form(real_user_name, user_dept)

            doc = Document()

            # 分別整理停車費 / 加油費明細，各自依日期排序
            parking_items = [
                {
                    "date": format_date_to_excel(p_item["date"]),
                    "title_type": "停車費",
                    "raw_bytes": p_item["raw_bytes"],
                    "file_ext": p_item["file_ext"],
                }
                for p_item in parking_receipts
            ]
            parking_items.sort(key=lambda x: x["date"])

            gas_items = [
                {
                    "date": format_date_to_excel(g_item["date"]),
                    "title_type": "加油費",
                    "raw_bytes": g_item["raw_bytes"],
                    "file_ext": g_item["file_ext"],
                }
                for g_item in gas_receipts
            ]
            gas_items.sort(key=lambda x: x["date"])

            # 版面採「停車費區塊」+「加油費區塊」各自用多欄表格整齊排版，
            # 取代舊版「每張收據獨自佔一整頁」的做法，節省紙張並方便核對。
            if parking_items:
                render_receipt_grid(
                    doc, "停車費明細", parking_items, items_per_row=2
                )

            if gas_items:
                # 區塊間的留白已經靠標題段落自己的 space_before 處理，
                # 這裡不再額外加空白段落，避免空白段落本身被單獨留在頁尾。
                render_receipt_grid(
                    doc, "加油費明細", gas_items, items_per_row=2
                )

            if not parking_items and not gas_items:
                doc.add_paragraph("尚無收據資料")

            # 檔名日期一律採台灣時區的「今天」，避免部署主機時區不同造成誤植
            output_date = get_taipei_now().strftime("%Y%m%d")
            # 檔名使用純中文真實全名
            word_filename = f"報支單據-{real_user_name}-{output_date}.docx"

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".docx"
            ) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as file:
                st.download_button(
                    label="\U0001F4E5 下載報支單據 (Word)",
                    data=file,
                    file_name=word_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

inject_enter_and_memory_js()
inject_custom_footer()
